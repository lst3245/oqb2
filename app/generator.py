"""
Word document generation module
"""
from flask import Blueprint, render_template, request, send_file, current_app, flash, redirect, url_for, session, jsonify, abort
from flask_login import login_required, current_user
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
from PIL import Image
import os
import re
import json
import logging
import shutil
import subprocess
import threading
import zipfile
import tempfile
from datetime import datetime
from collections import OrderedDict

logger = logging.getLogger(__name__)
from app import db
from app.models import Question, QuestionAsset, GeneratedFile
from app.utils import natural_sort, apply_multi_sort, SORT_FIELDS

generator_bp = Blueprint('generator', __name__, url_prefix='/generate')


def _require_generate_permission():
    """Check that current user can generate documents (not view-only). Aborts with 403 if not."""
    if not current_user.can_generate():
        abort(403)


@generator_bp.route('/', methods=['GET', 'POST'])
@login_required
def index():
    """Generation options page"""
    _require_generate_permission()
    # Accept question IDs from POST form data (preferred) or GET query params (fallback)
    filter_data = ''
    generation_options = {}
    
    if request.method == 'POST':
        question_ids = request.form.getlist('question_ids')
        filter_data = request.form.get('filter_data', '')
        # Store in session so page refreshes still work
        if question_ids:
            session['generator_question_ids'] = question_ids
            session['generator_filter_data'] = filter_data
            sort_config_str = request.form.get('sort_config')
            if sort_config_str:
                try:
                    session['sort_config'] = json.loads(sort_config_str)
                except json.JSONDecodeError:
                    pass
    else:
        question_ids = request.args.getlist('question_ids')
        if not question_ids:
            # Fallback to session (e.g. page refresh)
            question_ids = session.get('generator_question_ids', [])
        filter_data = session.get('generator_filter_data', '')
    
    # Check if regenerating from a saved file
    regen_file_id = request.args.get('regen_file_id') or request.form.get('regen_file_id')
    regen_display_name = ''
    is_regen = False
    # Track whether question IDs were explicitly provided via POST (user chose dashboard selection)
    has_explicit_qids = request.method == 'POST' and bool(request.form.getlist('question_ids'))
    if regen_file_id:
        try:
            gen_file = GeneratedFile.query.get(int(regen_file_id))
            if gen_file and (gen_file.user_id == current_user.id or current_user.is_super_admin):
                is_regen = True
                # Load generation options
                if gen_file.generation_options:
                    try:
                        generation_options = json.loads(gen_file.generation_options)
                    except (json.JSONDecodeError, TypeError):
                        pass
                # Load question IDs from saved options
                # Only override if user didn't explicitly provide IDs (e.g. chose dashboard selection)
                saved_qids = generation_options.get('question_ids', [])
                if saved_qids and not has_explicit_qids:
                    question_ids = [str(qid) for qid in saved_qids]
                    session['generator_question_ids'] = question_ids
                # Load filter data — always use the original file's filter, overriding any stale session value
                if gen_file.filter_data:
                    filter_data = gen_file.filter_data
                    session['generator_filter_data'] = filter_data
                # Load sort config
                saved_sort = generation_options.get('sort_config')
                if saved_sort:
                    try:
                        session['sort_config'] = json.loads(saved_sort) if isinstance(saved_sort, str) else saved_sort
                    except (json.JSONDecodeError, TypeError):
                        pass
                # Get display name for pre-filling
                if gen_file.display_name:
                    regen_display_name = re.sub(r'_\d{8}_\d{6}$', '', gen_file.display_name)
        except (ValueError, TypeError):
            pass
    
    if not question_ids:
        flash('No questions selected', 'warning')
        return redirect(url_for('dashboard.index'))
    
    # Get questions - preserve the selection order by using a dict
    questions_dict = {str(q.id): q for q in Question.query.filter(Question.id.in_(question_ids)).all()}
    questions = [questions_dict[qid] for qid in question_ids if qid in questions_dict]
    
    # Get sort config from session (from dashboard)
    sort_config = session.get('sort_config', [{"field": "qid", "direction": "asc"}])
    
    # Get available sort fields for the UI
    sort_fields = [{"value": key, "label": info["label"]} for key, info in SORT_FIELDS.items()]
    
    return render_template('generate.html', 
                          questions=questions,
                          question_ids=question_ids,
                          sort_config=sort_config,
                          sort_fields=sort_fields,
                          filter_data=filter_data,
                          generation_options=generation_options,
                          regen_display_name=regen_display_name,
                          is_regen=is_regen)

@generator_bp.route('/viewer', methods=['GET', 'POST'])
@login_required
def viewer():
    """Full-screen viewer/presentation mode page"""
    # Accept question IDs from POST form data (preferred) or GET query params (fallback)
    if request.method == 'POST':
        question_ids = request.form.getlist('question_ids')
        sort_config_str = request.form.get('sort_config')
        # Store in session so page refreshes still work
        if question_ids:
            session['viewer_question_ids'] = question_ids
        if sort_config_str:
            try:
                sort_config = json.loads(sort_config_str)
                session['viewer_sort_config'] = sort_config
            except json.JSONDecodeError:
                sort_config = session.get('sort_config', [{"field": "qid", "direction": "asc"}])
        else:
            sort_config = session.get('sort_config', [{"field": "qid", "direction": "asc"}])
    else:
        question_ids = request.args.getlist('question_ids')
        if not question_ids:
            # Fallback to session (e.g. page refresh)
            question_ids = session.get('viewer_question_ids', [])
        sort_config_str = request.args.get('sort_config')
        if sort_config_str:
            try:
                sort_config = json.loads(sort_config_str)
            except json.JSONDecodeError:
                sort_config = session.get('sort_config', [{"field": "qid", "direction": "asc"}])
        else:
            sort_config = session.get('viewer_sort_config', 
                         session.get('sort_config', [{"field": "qid", "direction": "asc"}]))
    
    if not question_ids:
        flash('No questions selected', 'warning')
        return redirect(url_for('dashboard.index'))
    
    # Get questions - preserve the selection order by using a dict
    questions_dict = {str(q.id): q for q in Question.query.filter(Question.id.in_(question_ids)).all()}
    questions_list = [questions_dict[qid] for qid in question_ids if qid in questions_dict]
    
    # Apply sort if custom sort mode
    questions_list = apply_multi_sort(questions_list, sort_config)
    
    # Prepare question data with assets
    questions_data = []
    for q in questions_list:
        # Get all assets for this question
        assets = QuestionAsset.query.filter_by(question_id=q.id).all()
        
        # Group assets by type
        que_assets = [a for a in assets if a.asset_type == 'QUE']
        ans_assets = [a for a in assets if a.asset_type == 'ANS']
        sol_assets = [a for a in assets if a.asset_type == 'SOL']
        
        questions_data.append({
            'id': q.id,
            'qid': q.qid,
            'year': q.year,
            'level': q.level,
            'q_type': q.q_type,
            'has_que': len(que_assets) > 0,
            'has_ans': len(ans_assets) > 0,
            'has_sol': len(sol_assets) > 0,
            'answer': q.answer,
            'comment': q.comment,
            'has_answer_text': bool(q.answer),
            'has_comment': bool(q.comment)
        })
    
    return render_template('viewer.html', 
                          questions=questions_data,
                          question_ids=question_ids)

@generator_bp.route('/api/viewer_asset/<int:question_id>/<asset_type>')
@login_required
def get_viewer_asset(question_id, asset_type):
    """Get asset URLs for viewer mode with language preference.
    Returns an array of parts for multi-image questions.
    """
    preferred_language = request.args.get('lang', 'EN')
    
    # Get all assets for this question and type
    assets = QuestionAsset.query.filter_by(
        question_id=question_id,
        asset_type=asset_type
    ).all()
    
    if not assets:
        # If SOL not found, try ANS and vice versa
        fallback_type = 'ANS' if asset_type == 'SOL' else 'SOL' if asset_type == 'ANS' else None
        if fallback_type:
            assets = QuestionAsset.query.filter_by(
                question_id=question_id,
                asset_type=fallback_type
            ).all()
            if assets:
                asset_type = fallback_type  # Update to indicate the fallback
    
    if not assets:
        return jsonify({'error': 'Asset not found', 'asset_type': asset_type}), 404
    
    # Sort by language preference and format. Priority: IMG > MD > DOC.
    def lang_order(asset):
        if asset.language == preferred_language:
            return 0
        elif asset.language == 'BI':
            return 1
        else:
            return 2

    _FMT_RANK = {'IMG': 0, 'MD': 1, 'DOC': 2}
    def format_order(asset):
        return _FMT_RANK.get(asset.file_format, 99)

    sorted_assets = sorted(assets, key=lambda a: (format_order(a), lang_order(a), a.part_number))

    # Pick the best language+format group, then return all parts for that group
    best = sorted_assets[0]
    selected = [a for a in sorted_assets
                if a.file_format == best.file_format and a.language == best.language]
    # Ensure ordered by part_number
    selected.sort(key=lambda a: a.part_number)

    parts = []
    for a in selected:
        parts.append({
            'id': a.id,
            'type': a.asset_type,
            'format': a.file_format,
            'language': a.language,
            'part_number': a.part_number,
            'url': f"/dashboard/files/{a.file_path}"
        })

    # Return parts array + backward-compat top-level fields from first part
    result = {
        'parts': parts,
        'id': parts[0]['id'],
        'type': parts[0]['type'],
        'format': parts[0]['format'],
        'language': parts[0]['language'],
        'url': parts[0]['url'],
        'asset_type': asset_type,
    }

    # For MD assets, also include the rendered HTML so the viewer can show
    # the markdown inline (no download required).
    if best.file_format == 'MD':
        from app import md_render
        source_path = current_app.config['SOURCE_PATH']
        abs_path = os.path.join(source_path, *best.file_path.split('/'))
        result['html'] = md_render.render_file(best.id, abs_path)

    return jsonify(result)

@generator_bp.route('/create', methods=['POST'])
@login_required
def create_document():
    """Start background generation of Word document from selected questions"""
    _require_generate_permission()
    
    # Get parameters
    question_ids = request.form.getlist('question_ids')
    sort_mode = request.form.get('sort_mode', 'custom')
    sort_config_str = request.form.get('sort_config', '')
    answer_mode = request.form.get('answer_mode', 'QUE_ONLY')
    display_name = request.form.get('display_name', '').strip()
    filter_data = request.form.get('filter_data', '')
    
    # MC spacing settings
    mc_before_mode = request.form.get('mc_before_mode', 'lines')
    mc_before_lines = int(request.form.get('mc_before_lines', 0))
    mc_after_mode = request.form.get('mc_after_mode', 'lines')
    mc_after_lines = int(request.form.get('mc_after_lines', 1))
    
    # CQ spacing settings
    cq_before_mode = request.form.get('cq_before_mode', 'page')
    cq_before_lines = int(request.form.get('cq_before_lines', 0))
    cq_after_mode = request.form.get('cq_after_mode', 'page')
    cq_after_lines = int(request.form.get('cq_after_lines', 0))
    
    # Show QID options
    show_qid = request.form.get('show_qid') == 'on'
    show_qid_answer = request.form.get('show_qid_answer') == 'on'
    show_correct_pct = request.form.get('show_correct_pct') == 'on'
    show_seq_no = request.form.get('show_seq_no') == 'on'
    seq_start = max(1, int(request.form.get('seq_start', 1) or 1))
    show_page_no = request.form.get('show_page_no') == 'on'
    keep_together = request.form.get('keep_together') == 'on'
    apply_spacing_to_ans = request.form.get('apply_spacing_to_ans') == 'on'
    denote_cross_topic = request.form.get('denote_cross_topic') == 'on'
    
    # Topic/Chapter display options
    info_fields = {
        'topic': request.form.get('info_topic') == 'on',
        'subtopic': request.form.get('info_subtopic') == 'on',
        'chapter': request.form.get('info_chapter') == 'on',
        'subchapter': request.form.get('info_subchapter') == 'on',
    }
    section_fields = {
        'topic': request.form.get('section_topic') == 'on',
        'subtopic': request.form.get('section_subtopic') == 'on',
        'chapter': request.form.get('section_chapter') == 'on',
        'subchapter': request.form.get('section_subchapter') == 'on',
    }
    split_fields = {
        'topic': request.form.get('split_topic') == 'on',
        'subtopic': request.form.get('split_subtopic') == 'on',
        'chapter': request.form.get('split_chapter') == 'on',
        'subchapter': request.form.get('split_subchapter') == 'on',
    }
    
    preferred_language = request.form.get('preferred_language', 'EN')
    answer_preference = request.form.get('answer_preference', 'image_first')
    format_priority = _parse_format_priority(request.form.get('format_priority', ''))

    if not question_ids:
        return jsonify({'error': 'No questions selected'}), 400

    # Filter out questions from view-only subjects
    if not current_user.is_super_admin:
        subject_roles = current_user.get_subject_roles()
        allowed_subjects = {sid for sid, role in subject_roles.items() if role in ('user', 'admin')}
        qs = Question.query.filter(Question.id.in_(question_ids)).all()
        question_ids = [str(q.id) for q in qs if q.subject in allowed_subjects]
        if not question_ids:
            return jsonify({'error': 'No questions available for generation with your permissions'}), 403

    # Build generation options for storage
    generation_options = {
        'sort_mode': sort_mode,
        'sort_config': sort_config_str,
        'answer_mode': answer_mode,
        'mc_before_mode': mc_before_mode, 'mc_before_lines': mc_before_lines,
        'mc_after_mode': mc_after_mode, 'mc_after_lines': mc_after_lines,
        'cq_before_mode': cq_before_mode, 'cq_before_lines': cq_before_lines,
        'cq_after_mode': cq_after_mode, 'cq_after_lines': cq_after_lines,
        'show_qid': show_qid, 'show_qid_answer': show_qid_answer,
        'show_correct_pct': show_correct_pct,
        'show_seq_no': show_seq_no, 'seq_start': seq_start, 'show_page_no': show_page_no,
        'keep_together': keep_together,
        'apply_spacing_to_ans': apply_spacing_to_ans,
        'denote_cross_topic': denote_cross_topic,
        'info_fields': info_fields, 'section_fields': section_fields,
        'split_fields': split_fields,
        'preferred_language': preferred_language,
        'answer_preference': answer_preference,
        'format_priority': ','.join(format_priority),
        'question_ids': question_ids,
    }
    
    # Create filename
    any_split = any(split_fields.values())
    file_ext = '.zip' if any_split else '.docx'
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    if not display_name:
        display_name = f'questions_{timestamp}'
        filename = f'{display_name}{file_ext}'
    else:
        filename = f'{display_name}_{timestamp}{file_ext}'
    # Sanitize filename
    filename = "".join(c for c in filename if c.isalnum() or c in '._- ').strip()
    if not filename.endswith(file_ext):
        filename += file_ext
    
    # Create GeneratedFile record
    gen_file = GeneratedFile(
        user_id=current_user.id,
        display_name=display_name,
        filename=filename,
        status='pending',
        filter_data=filter_data if filter_data else None,
        generation_options=json.dumps(generation_options),
        question_count=len(question_ids),
    )
    db.session.add(gen_file)
    db.session.commit()
    gen_file_id = gen_file.id
    
    # Build spacing config
    spacing_config = {
        'mc': {
            'before_mode': mc_before_mode, 'before_lines': mc_before_lines,
            'after_mode': mc_after_mode, 'after_lines': mc_after_lines
        },
        'cq': {
            'before_mode': cq_before_mode, 'before_lines': cq_before_lines,
            'after_mode': cq_after_mode, 'after_lines': cq_after_lines
        }
    }
    
    # Spawn background thread
    app = current_app._get_current_object()
    thread = threading.Thread(
        target=_generate_in_background,
        args=(app, gen_file_id, question_ids, sort_mode, sort_config_str,
              answer_mode, spacing_config, show_qid, show_qid_answer,
              preferred_language, show_correct_pct, answer_preference,
              show_seq_no, seq_start, show_page_no, keep_together,
              apply_spacing_to_ans, denote_cross_topic,
              info_fields, section_fields, split_fields, filename,
              format_priority)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({'id': gen_file_id, 'status': 'pending', 'filename': filename})


def _generate_in_background(app, gen_file_id, question_ids, sort_mode, sort_config_str,
                            answer_mode, spacing_config, show_qid, show_qid_answer,
                            preferred_language, show_correct_pct, answer_preference,
                            show_seq_no, seq_start, show_page_no, keep_together,
                            apply_spacing_to_ans, denote_cross_topic,
                            info_fields, section_fields, split_fields, filename,
                            format_priority=None):
    """Background thread function to generate the Word document(s)"""
    format_priority = format_priority or list(_DEFAULT_FORMAT_PRIORITY)
    with app.app_context():
        gen_file = GeneratedFile.query.get(gen_file_id)
        if not gen_file:
            return
        
        gen_file.status = 'generating'
        db.session.commit()
        
        try:
            # Get questions
            questions_dict = {str(q.id): q for q in Question.query.filter(Question.id.in_(question_ids)).all()}
            
            if not questions_dict:
                gen_file.status = 'failed'
                gen_file.error_message = 'No valid questions found'
                db.session.commit()
                return
            
            # Sort questions
            if sort_mode == 'selection':
                questions = [questions_dict[qid] for qid in question_ids if qid in questions_dict]
            else:
                try:
                    sort_config = json.loads(sort_config_str) if sort_config_str else [{"field": "qid", "direction": "asc"}]
                except json.JSONDecodeError:
                    sort_config = [{"field": "qid", "direction": "asc"}]
                questions = list(questions_dict.values())
                questions = apply_multi_sort(questions, sort_config)
            
            any_split = split_fields and any(split_fields.values())
            output_path = app.config['OUTPUT_PATH']
            filepath = os.path.join(output_path, filename)
            
            if any_split:
                # Split questions into groups based on split_fields
                groups = _split_questions_into_groups(questions, split_fields)
                
                # Generate one doc per group and zip them
                with tempfile.TemporaryDirectory() as tmpdir:
                    docx_files = []  # (filename, filepath) tuples
                    used_names = set()  # track names to avoid duplicates
                    
                    for group_label, group_questions in groups.items():
                        doc = create_word_document(
                            group_questions, answer_mode, spacing_config,
                            show_qid, show_qid_answer, preferred_language,
                            show_correct_pct, answer_preference,
                            show_seq_no, seq_start, show_page_no, keep_together,
                            info_fields, section_fields,
                            apply_spacing_to_ans=apply_spacing_to_ans,
                            denote_cross_topic=denote_cross_topic,
                            format_priority=format_priority,
                        )
                        # Build per-file name from the group label, dedup if needed
                        safe_label = _sanitize_filename(group_label)
                        docx_name = f'{safe_label}.docx'
                        counter = 2
                        while docx_name in used_names:
                            docx_name = f'{safe_label} ({counter}).docx'
                            counter += 1
                        used_names.add(docx_name)
                        
                        docx_path = os.path.join(tmpdir, docx_name)
                        doc.save(docx_path)
                        docx_files.append((docx_name, docx_path))
                    
                    # Create zip
                    with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for docx_name, docx_path in docx_files:
                            zf.write(docx_path, docx_name)
            else:
                # Single document
                doc = create_word_document(
                    questions, answer_mode, spacing_config,
                    show_qid, show_qid_answer, preferred_language,
                    show_correct_pct, answer_preference,
                    show_seq_no, seq_start, show_page_no, keep_together,
                    info_fields, section_fields,
                    apply_spacing_to_ans=apply_spacing_to_ans,
                    denote_cross_topic=denote_cross_topic,
                    format_priority=format_priority,
                )
                doc.save(filepath)
            
            gen_file.status = 'completed'
            gen_file.completed_at = datetime.utcnow()
            db.session.commit()
            
        except Exception as e:
            gen_file.status = 'failed'
            gen_file.error_message = str(e)
            db.session.commit()


def _split_questions_into_groups(questions, split_fields):
    """
    Split questions into ordered groups based on split_fields.
    Returns an OrderedDict of {group_label: [questions]}.
    The group label is built from the field values (e.g. "Topic - Subtopic").
    """
    groups = OrderedDict()
    
    for question in questions:
        key = _get_section_key(question, split_fields)
        
        # Build human-readable label from key parts
        label = _build_split_label(key)
        
        if label not in groups:
            groups[label] = []
        groups[label].append(question)
    
    return groups


def _build_split_label(key):
    """Build a human-readable label from a section key tuple for split filenames."""
    topic_parts = []
    chapter_parts = []
    
    for field_name, value in key:
        val = value or 'Unknown'
        if field_name in ('topic', 'subtopic'):
            topic_parts.append(val)
        elif field_name in ('chapter', 'subchapter'):
            chapter_parts.append(val)
    
    parts = []
    if topic_parts:
        parts.append(' - '.join(topic_parts))
    if chapter_parts:
        parts.append(' - '.join(chapter_parts))
    
    return ' _ '.join(parts) if parts else 'Uncategorized'


def _sanitize_filename(name):
    """Sanitize a string for use as a filename."""
    # Replace characters not safe for filenames
    safe = "".join(c for c in name if c.isalnum() or c in '._- ').strip()
    return safe if safe else 'Untitled'


@generator_bp.route('/status/<int:file_id>')
@login_required
def generation_status(file_id):
    """Check generation status (for polling)"""
    gen_file = GeneratedFile.query.get_or_404(file_id)
    
    # Only owner or super admin
    if gen_file.user_id != current_user.id and not current_user.is_super_admin:
        return jsonify({'error': 'Access denied'}), 403
    
    return jsonify({
        'id': gen_file.id,
        'status': gen_file.status,
        'error_message': gen_file.error_message,
        'display_name': gen_file.display_name,
        'filename': gen_file.filename,
    })


@generator_bp.route('/download/<int:file_id>')
@login_required
def download_file(file_id):
    """Download a completed generated file"""
    _require_generate_permission()
    gen_file = GeneratedFile.query.get_or_404(file_id)
    
    # Only owner or super admin
    if gen_file.user_id != current_user.id and not current_user.is_super_admin:
        flash('Access denied', 'danger')
        return redirect(url_for('user.files'))
    
    if gen_file.status != 'completed':
        flash('File is not ready for download', 'warning')
        return redirect(url_for('user.files'))
    
    output_path = current_app.config['OUTPUT_PATH']
    filepath = os.path.join(output_path, gen_file.filename)
    
    if not os.path.exists(filepath):
        flash('File not found on disk', 'danger')
        return redirect(url_for('user.files'))
    
    # Determine mimetype based on file extension
    if gen_file.filename.endswith('.zip'):
        mimetype = 'application/zip'
    else:
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    return send_file(
        filepath,
        as_attachment=True,
        download_name=gen_file.filename,
        mimetype=mimetype
    )

def get_question_spacing_config(question, spacing_config):
    """
    Get spacing settings for a question based on its type (MC or CQ)
    Returns dict with before_mode, before_lines, after_mode, after_lines
    """
    q_type = question.q_type.upper() if question.q_type else 'CQ'  # Default to CQ if not set
    
    if q_type == 'MC':
        return spacing_config['mc']
    else:
        # CQ or any other type uses CQ settings
        return spacing_config['cq']


def add_before_spacing(doc, spacing, last_had_page_break, is_first):
    """
    Add spacing before a question based on settings.
    Returns whether we effectively have a page break before this question.
    
    Smart logic: If previous question added page break, don't add another one
    even if "start from new page" is selected.
    """
    if is_first:
        # First question - no spacing before
        return False
    
    if spacing['before_mode'] == 'page':
        # Want to start from new page
        if not last_had_page_break:
            # Previous question didn't add page break, so we need one
            doc.add_page_break()
        # Either way, we're now at a new page
        return True
    else:
        # Skip lines before
        for _ in range(spacing['before_lines']):
            doc.add_paragraph()
        return False


def add_after_spacing(doc, spacing):
    """
    Add spacing after a question based on settings.
    Returns whether we added a page break.
    """
    if spacing['after_mode'] == 'page':
        doc.add_page_break()
        return True
    else:
        # Skip lines after
        for _ in range(spacing['after_lines']):
            doc.add_paragraph()
        return False


def _define_oqb_styles(doc):
    """
    Create (or ensure) the four OQB custom paragraph styles used throughout the document.
    Defining them here keeps the visual defaults in one place and allows users to
    override them inside Word after generation.

    Styles created
    ──────────────
    OQB Section Heading  → centred, bold, 14 pt  (topic/chapter change headings)
    OQB Question ID      → bold, 12 pt            (question number / QID / % line)
    OQB Question Info    → italic, 10 pt, grey     (topic-chapter info line)
    OQB Body Text        → 11 pt                   (answer text fallbacks, placeholders)
    """
    styles = doc.styles

    def _make(name, font_size_pt, bold=False, italic=False,
              align=None, colour_rgb=None, base_style='Normal'):
        # Re-use if the style already exists (e.g. from a template)
        if name in [s.name for s in styles]:
            return styles[name]
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base_style]
        style.quick_style = True           # appear in the Styles gallery in Word
        font = style.font
        font.size = Pt(font_size_pt)
        font.bold = bold
        font.italic = italic
        if colour_rgb:
            font.color.rgb = colour_rgb
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if align:
            pf.alignment = align
        return style

    _make('OQB Section Heading', 14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _make('OQB Question ID',     12, bold=True)
    _make('OQB Question Info',   10, italic=True,
          colour_rgb=RGBColor(100, 100, 100))
    _make('OQB Body Text',       11)


def _add_page_numbers(doc):
    """Add auto-generated page numbers at the bottom centre of every page."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Build the PAGE field XML: renders as the current page number
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_end)
        
        run.font.size = Pt(10)


def _get_section_key(question, section_fields):
    """Get the current section key tuple based on which section fields are enabled."""
    key = []
    if section_fields.get('topic'):
        key.append(('topic', question.major_topic.name if question.major_topic else None))
    if section_fields.get('subtopic'):
        key.append(('subtopic', question.major_subtopic.name if question.major_subtopic else None))
    if section_fields.get('chapter'):
        key.append(('chapter', question.chapter.name if question.chapter else None))
    if section_fields.get('subchapter'):
        key.append(('subchapter', question.subchapter.name if question.subchapter else None))
    return tuple(key)


def _add_section_heading(doc, question, prev_key, section_fields, keep_together=False):
    """
    If any tracked section field changed, insert a centered bold heading.
    Returns the new key.
    """
    current_key = _get_section_key(question, section_fields)
    if current_key == prev_key:
        return prev_key  # no change
    
    # Build heading text from the parts that are enabled
    parts = []
    topic_part = []
    chapter_part = []
    
    if section_fields.get('topic') and question.major_topic:
        topic_part.append(question.major_topic.name)
    if section_fields.get('subtopic') and question.major_subtopic:
        topic_part.append(question.major_subtopic.name)
    if section_fields.get('chapter') and question.chapter:
        chapter_part.append(question.chapter.name)
    if section_fields.get('subchapter') and question.subchapter:
        chapter_part.append(question.subchapter.name)
    
    if topic_part:
        parts.append(' - '.join(topic_part))
    if chapter_part:
        parts.append(' - '.join(chapter_part))
    
    heading_text = ' | '.join(parts) if parts else None
    
    if heading_text:
        heading = doc.add_paragraph(style='OQB Section Heading')
        heading.add_run(heading_text)
        if keep_together:
            heading.paragraph_format.keep_with_next = True
    
    return current_key


_DPI = 96  # assumed screen DPI for image size conversion


# Default format priority for picking which asset to render when a question
# has multiple file_format options (IMG, MD, DOC). Lower index = preferred.
_DEFAULT_FORMAT_PRIORITY = ('IMG', 'MD', 'DOC')


def _parse_format_priority(raw):
    """Parse a comma-separated `format_priority` form value into a 3-element
    list. Defaults to ('IMG', 'MD', 'DOC') and tolerates missing / unknown
    tokens by appending the defaults in order."""
    valid = set(_DEFAULT_FORMAT_PRIORITY)
    parts = [p.strip().upper() for p in (raw or '').split(',') if p.strip()]
    out = []
    for p in parts:
        if p in valid and p not in out:
            out.append(p)
    for p in _DEFAULT_FORMAT_PRIORITY:
        if p not in out:
            out.append(p)
    return out[:len(_DEFAULT_FORMAT_PRIORITY)]


# Match $$...$$ blocks (display math), greedy across newlines but non-greedy
# between two consecutive `$$` so adjacent blocks don't merge.
_MD_DISPLAY_MATH_RE = re.compile(r'\$\$([\s\S]+?)\$\$', re.MULTILINE)


def _preprocess_md_for_pandoc(src):
    """
    Normalise a user-authored .md file so pandoc renders it the way the
    in-browser preview does.

    1. Collapse blank lines INSIDE $$...$$ display-math blocks. Pandoc's
       `tex_math_dollars` extension treats a blank line as a paragraph break,
       so `$$\\n\\nformula\\n\\n$$` becomes three separate <p> elements and the
       math never renders. The live editor doesn't care (we KaTeX-render the
       blocks ourselves), but pandoc does, so we squash the blank lines here.
    """
    def _collapse_blank_lines(match):
        inner = match.group(1)
        # Replace any run of (newline + whitespace-only line(s) + newline) with
        # a single newline so the math content is one contiguous paragraph.
        inner = re.sub(r'\n[ \t]*\n+', '\n', inner)
        return '$$' + inner + '$$'
    return _MD_DISPLAY_MATH_RE.sub(_collapse_blank_lines, src)


def _append_md_via_pandoc(master_doc, md_abs_path):
    """
    Convert a single self-contained Markdown file to .docx via pandoc, then
    splice it into `master_doc` using docxcompose. Raises on any failure so
    the caller can fall back to a placeholder.

    The Markdown file is expected to be self-contained (LaTeX math via
    `$...$` / `$$...$$` and base64-embedded images). Pandoc converts dollar
    math to native Word OMML equations.

    Pre-processing notes:
      * Blank lines inside $$...$$ are collapsed (see _preprocess_md_for_pandoc).
      * `-implicit_figures` is passed to pandoc's reader so standalone images
        do NOT get rendered as captioned figures with the alt text as caption.
    """
    pandoc = current_app.config.get('PANDOC_PATH', 'pandoc')
    # Resolve to a usable absolute path.
    # shutil.which only looks at the PATH that was set when the process *started*,
    # so on Windows a freshly-installed pandoc is often invisible here even though
    # it works fine in a new terminal. Fall through to well-known install locations.
    pandoc_bin = shutil.which(pandoc)
    if not pandoc_bin and pandoc == 'pandoc':
        _WIN_PANDOC_CANDIDATES = [
            r'C:\Program Files\Pandoc\pandoc.exe',
            r'C:\Program Files (x86)\Pandoc\pandoc.exe',
            os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Pandoc', 'pandoc.exe'),
            os.path.join(os.environ.get('APPDATA', ''), 'Pandoc', 'pandoc.exe'),
        ]
        for candidate in _WIN_PANDOC_CANDIDATES:
            if candidate and os.path.isfile(candidate):
                pandoc_bin = candidate
                break
    pandoc_bin = pandoc_bin or pandoc

    # Read the source and pre-process. We write to a tmp file because pandoc
    # accepts a path more reliably than stdin on Windows (no encoding fuss).
    try:
        with open(md_abs_path, 'r', encoding='utf-8') as src_f:
            src = src_f.read()
    except OSError as e:
        raise RuntimeError(f'failed to read md file: {e}') from e
    except UnicodeDecodeError as e:
        raise RuntimeError(
            f'md file is not valid UTF-8 ({e.reason} at byte {e.start})'
        ) from e
    src = _preprocess_md_for_pandoc(src)

    with tempfile.TemporaryDirectory(prefix='oqb_md_') as tmpdir:
        prepped_md = os.path.join(tmpdir, 'input.md')
        with open(prepped_md, 'w', encoding='utf-8') as out_f:
            out_f.write(src)

        out_docx = os.path.join(tmpdir, 'fragment.docx')
        cmd = [
            pandoc_bin,
            '--from=markdown+tex_math_dollars+tex_math_double_backslash-implicit_figures',
            '--to=docx',
            '-o', out_docx,
            prepped_md,
        ]
        try:
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
                check=False,
            )
        except FileNotFoundError as e:
            raise RuntimeError(
                f'pandoc not found at {pandoc!r}. Install pandoc or set PANDOC_PATH.'
            ) from e
        except subprocess.TimeoutExpired as e:
            raise RuntimeError('pandoc timed out converting Markdown') from e

        if result.returncode != 0:
            err = (result.stderr or b'').decode('utf-8', errors='replace').strip()
            raise RuntimeError(f'pandoc failed (rc={result.returncode}): {err[:400]}')

        if not os.path.exists(out_docx):
            raise RuntimeError('pandoc produced no output file')

        try:
            fragment = Document(out_docx)
        except Exception as e:
            raise RuntimeError(f'failed to open pandoc output: {e}') from e

        # docxcompose appends the fragment's body elements into the master,
        # remapping styles, numbering, and media references.
        Composer(master_doc).append(fragment)


def create_word_document(questions, answer_mode, spacing_config, show_qid, show_qid_answer, preferred_language='EN', show_correct_pct=False, answer_preference='image_first', show_seq_no=False, seq_start=1, show_page_no=False, keep_together=False, info_fields=None, section_fields=None, apply_spacing_to_ans=False, denote_cross_topic=False, format_priority=None):
    """
    Create Word document with questions.

    Args:
        questions: List of Question objects
        answer_mode: One of QUE_ONLY, QUE_ANS, QUE_SOL, QUE_THEN_ANS, QUE_THEN_SOL
        spacing_config: Dict with mc and cq sub-dicts containing before_mode, before_lines, after_mode, after_lines
        show_qid: Show question ID for questions
        show_qid_answer: Show question ID for answers/solutions
        preferred_language: 'EN' or 'CH' - order: preferred > BI > other
        show_correct_pct: Show correct percentage with question ID (format: "QID [X%]")
        answer_preference: 'image_first' or 'text_first' - for ANS content, prefer image or text
        show_seq_no: Show sequential question number (1. 2. 3. ...)
        seq_start: Starting number for sequential numbering (default 1)
        show_page_no: Show page numbers at bottom centre
        keep_together: Keep question info with image (prevent page split)
        info_fields: Dict of bools for per-question info line (topic/subtopic/chapter/subchapter)
        section_fields: Dict of bools for section heading on change (topic/subtopic/chapter/subchapter)
        apply_spacing_to_ans: If True, apply the same MC/CQ spacing to ANS/SOL in THEN modes.
                              If False (default), use minimal spacing (1 line between items).
        denote_cross_topic: If True and a question has minor topics, force info line to show
                            major topic with "[Cross Topic: ...]" annotation.
        format_priority: Ordered list (or tuple) of file_format strings — subset
                         of {'IMG','MD','DOC'}. When a question has multiple
                         available formats for the same asset type/language,
                         the first format in this list wins. Defaults to
                         ('IMG','MD','DOC').
    """
    if not format_priority:
        format_priority = list(_DEFAULT_FORMAT_PRIORITY)
    if info_fields is None:
        info_fields = {}
    if section_fields is None:
        section_fields = {}
    
    any_section_heading = any(section_fields.values())
    
    # Minimal spacing for ANS/SOL when apply_spacing_to_ans is False:
    # 0 lines before, 1 line after, no page breaks.
    minimal_ans_spacing = {
        'before_mode': 'lines', 'before_lines': 0,
        'after_mode': 'lines', 'after_lines': 1,
    }
    
    doc = Document()
    
    # Define OQB custom styles (must be done before any content is added)
    _define_oqb_styles(doc)
    
    # Set page size to A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 height
    section.page_width = Cm(21.0)   # A4 width
    
    # Set narrow margins (0.5 inches = 1.27 cm)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    
    # Add page numbers if requested
    if show_page_no:
        _add_page_numbers(doc)
    
    source_path = current_app.config['SOURCE_PATH']
    
    # Track if last question ended with a page break
    last_had_page_break = False
    # Track section heading key for change detection
    prev_section_key = None
    
    # Answer modes:
    # QUE_ONLY - questions only
    # QUE_ANS - question followed by answer
    # QUE_SOL - question followed by solution
    # QUE_THEN_ANS - all questions first, then all answers
    # QUE_THEN_SOL - all questions first, then all solutions
    
    if answer_mode == 'QUE_THEN_ANS':
        # Add all questions first
        for i, question in enumerate(questions):
            spacing = get_question_spacing_config(question, spacing_config)
            
            # Section heading on change (only for QUE section)
            if any_section_heading:
                prev_section_key = _add_section_heading(doc, question, prev_section_key, section_fields, keep_together)
            
            had_pb = add_before_spacing(doc, spacing, last_had_page_break, i == 0)
            seq_no = (seq_start + i) if show_seq_no else None
            add_question_content_to_doc(doc, question, 'QUE', show_qid, source_path, preferred_language, show_correct_pct, seq_no=seq_no, info_fields=info_fields, keep_together=keep_together, denote_cross_topic=denote_cross_topic, format_priority=format_priority)
            last_had_page_break = add_after_spacing(doc, spacing)
        
        # Then add all answers - always start on new page
        doc.add_page_break()
        heading = doc.add_paragraph('ANSWERS', style='OQB Section Heading')
        doc.add_paragraph()
        last_had_page_break = False
        
        for i, question in enumerate(questions):
            spacing = get_question_spacing_config(question, spacing_config) if apply_spacing_to_ans else minimal_ans_spacing
            add_before_spacing(doc, spacing, last_had_page_break, i == 0)
            seq_no = (seq_start + i) if show_seq_no else None
            add_question_content_to_doc(doc, question, 'ANS', show_qid_answer, source_path, preferred_language, show_correct_pct, answer_preference, seq_no=seq_no, keep_together=keep_together, denote_cross_topic=denote_cross_topic, format_priority=format_priority)
            last_had_page_break = add_after_spacing(doc, spacing)
    
    elif answer_mode == 'QUE_THEN_SOL':
        # Add all questions first
        for i, question in enumerate(questions):
            spacing = get_question_spacing_config(question, spacing_config)
            
            # Section heading on change (only for QUE section)
            if any_section_heading:
                prev_section_key = _add_section_heading(doc, question, prev_section_key, section_fields, keep_together)
            
            add_before_spacing(doc, spacing, last_had_page_break, i == 0)
            seq_no = (seq_start + i) if show_seq_no else None
            add_question_content_to_doc(doc, question, 'QUE', show_qid, source_path, preferred_language, show_correct_pct, seq_no=seq_no, info_fields=info_fields, keep_together=keep_together, denote_cross_topic=denote_cross_topic, format_priority=format_priority)
            last_had_page_break = add_after_spacing(doc, spacing)
        
        # Then add all solutions - always start on new page
        doc.add_page_break()
        heading = doc.add_paragraph('SOLUTIONS', style='OQB Section Heading')
        doc.add_paragraph()
        last_had_page_break = False
        
        for i, question in enumerate(questions):
            spacing = get_question_spacing_config(question, spacing_config) if apply_spacing_to_ans else minimal_ans_spacing
            add_before_spacing(doc, spacing, last_had_page_break, i == 0)
            seq_no = (seq_start + i) if show_seq_no else None
            add_question_content_to_doc(doc, question, 'SOL', show_qid_answer, source_path, preferred_language, show_correct_pct, seq_no=seq_no, keep_together=keep_together, denote_cross_topic=denote_cross_topic, format_priority=format_priority)
            last_had_page_break = add_after_spacing(doc, spacing)
    
    else:
        # Add questions with optional answers/solutions
        for i, question in enumerate(questions):
            spacing = get_question_spacing_config(question, spacing_config)
            
            # Section heading on change
            if any_section_heading:
                prev_section_key = _add_section_heading(doc, question, prev_section_key, section_fields, keep_together)
            
            add_before_spacing(doc, spacing, last_had_page_break, i == 0)
            seq_no = (seq_start + i) if show_seq_no else None
            add_question_content_to_doc(doc, question, 'QUE', show_qid, source_path, preferred_language, show_correct_pct, seq_no=seq_no, info_fields=info_fields, keep_together=keep_together, denote_cross_topic=denote_cross_topic, format_priority=format_priority)
            
            # Add answer/solution if requested (no extra spacing between Q and A/S)
            if answer_mode == 'QUE_ANS':
                add_question_content_to_doc(doc, question, 'ANS', show_qid_answer, source_path, preferred_language, show_correct_pct, answer_preference, keep_together=keep_together, format_priority=format_priority)
            elif answer_mode == 'QUE_SOL':
                add_question_content_to_doc(doc, question, 'SOL', show_qid_answer, source_path, preferred_language, show_correct_pct, keep_together=keep_together, format_priority=format_priority)
            
            last_had_page_break = add_after_spacing(doc, spacing)
    
    return doc

def add_question_content_to_doc(doc, question, asset_type, show_qid, source_path, preferred_language='EN', show_correct_pct=False, answer_preference='image_first', seq_no=None, info_fields=None, keep_together=False, denote_cross_topic=False, format_priority=None):
    """
    Add a question (or answer/solution) content to the document.
    Spacing is handled separately by add_before_spacing and add_after_spacing.
    
    Args:
        preferred_language: 'EN' or 'CH' - order: preferred > BI > other
        show_correct_pct: Show correct percentage (format: "QID [X%]" or just "[X%]" if no QID)
                          Only shown for QUE type, not for ANS or SOL
        answer_preference: 'image_first' or 'text_first' - for ANS content, prefer image or text
        seq_no: Sequential question number (int) or None to skip
        info_fields: Dict of bools for per-question info line (topic/subtopic/chapter/subchapter)
        keep_together: Set keep_with_next on heading/info paragraphs
        denote_cross_topic: If True and a question has minor topics, force info line to show
                            major topic with "[Cross Topic: ...]" annotation.
    """
    if info_fields is None:
        info_fields = {}
    
    # Build heading: "{seq_no}. {QID} [{pct}%]" — each part optional
    has_seq = seq_no is not None
    has_qid = show_qid
    has_pct = show_correct_pct and asset_type == 'QUE' and question.correct_percentage is not None
    
    if has_seq or has_qid or has_pct:
        parts = []
        if has_seq:
            parts.append(f"{seq_no}.")
        if has_qid:
            parts.append(question.qid)
        if has_pct:
            parts.append(f"[{question.correct_percentage}%]")
        
        heading_text = " ".join(parts)
        if heading_text:
            heading = doc.add_paragraph(style='OQB Question ID')
            heading.add_run(heading_text)
            if keep_together:
                heading.paragraph_format.keep_with_next = True
    
    # Add per-question info line (QUE only): "Topic - Subtopic | Chapter - Subchapter"
    # If denote_cross_topic is on AND question has minor topics, force the info line
    # to show the major topic with "[Cross Topic: <minor topics>]" suffix.
    if asset_type == 'QUE':
        # Determine if this question is cross-topic (has minor topics)
        minor_topic_names = []
        if denote_cross_topic:
            try:
                minor_topic_names = [t.name for t in question.minor_topics]
            except Exception:
                minor_topic_names = []
        is_cross_topic = denote_cross_topic and bool(minor_topic_names)
        
        # Show info line if any info_field is enabled OR cross-topic is being denoted
        if any(info_fields.values()) or is_cross_topic:
            topic_part = []
            chapter_part = []
            
            # If cross-topic is being denoted, force major topic to be shown
            show_major_topic = info_fields.get('topic') or is_cross_topic
            
            if show_major_topic and question.major_topic:
                topic_part.append(question.major_topic.name)
            if info_fields.get('subtopic') and question.major_subtopic:
                topic_part.append(question.major_subtopic.name)
            if info_fields.get('chapter') and question.chapter:
                chapter_part.append(question.chapter.name)
            if info_fields.get('subchapter') and question.subchapter:
                chapter_part.append(question.subchapter.name)
            
            info_parts = []
            if topic_part:
                topic_str = ' - '.join(topic_part)
                # Append cross-topic annotation if applicable
                if is_cross_topic:
                    topic_str += f' [Cross Topic: {", ".join(minor_topic_names)}]'
                info_parts.append(topic_str)
            elif is_cross_topic:
                # No topic_part but still cross-topic — show the annotation alone
                info_parts.append(f'[Cross Topic: {", ".join(minor_topic_names)}]')
            if chapter_part:
                info_parts.append(' - '.join(chapter_part))
            
            info_text = ' | '.join(info_parts)
            if info_text:
                info_para = doc.add_paragraph(style='OQB Question Info')
                info_para.add_run(info_text)
                if keep_together:
                    info_para.paragraph_format.keep_with_next = True
    
    # For ANS type, handle answer_preference (text_first vs image_first)
    if asset_type == 'ANS' and answer_preference == 'text_first':
        # Text first: use answer text if available, fall back to image
        if question.answer:
            para = doc.add_paragraph(style='OQB Body Text')
            para.add_run(question.answer)
            return
        # Fall through to image if no text
    
    # Get all available assets for this question and type
    assets = QuestionAsset.query.filter_by(
        question_id=question.id,
        asset_type=asset_type
    ).all()
    
    if not assets:
        # No image asset found
        if asset_type == 'ANS' and answer_preference == 'image_first' and question.answer:
            # Image first but no image available — fall back to answer text
            para = doc.add_paragraph(style='OQB Body Text')
            para.add_run(question.answer)
            return
        # No asset found, add placeholder
        para = doc.add_paragraph(style='OQB Body Text')
        run = para.add_run(f'[{asset_type} not available for {question.qid}]')
        run.italic = True
        return
    
    # Sort assets by language preference and format
    # Language order: preferred > BI > other
    def lang_order(asset):
        if asset.language == preferred_language:
            return 0
        elif asset.language == 'BI':
            return 1
        else:
            return 2

    # Format order: caller-supplied priority (default IMG > MD > DOC).
    _fp = format_priority or list(_DEFAULT_FORMAT_PRIORITY)
    _FMT_RANK = {fmt: i for i, fmt in enumerate(_fp)}
    def format_order(asset):
        return _FMT_RANK.get(asset.file_format, 99)

    # Sort by format first (per priority), then by language, then by part_number
    sorted_assets = sorted(assets, key=lambda a: (format_order(a), lang_order(a), a.part_number))

    if not sorted_assets:
        # No asset found, add placeholder
        para = doc.add_paragraph(style='OQB Body Text')
        run = para.add_run(f'[{asset_type} not available for {question.qid}]')
        run.italic = True
        return

    # Pick the best language+format group, then include all parts for that group
    best = sorted_assets[0]
    selected_assets = [a for a in sorted_assets
                       if a.file_format == best.file_format and a.language == best.language]
    selected_assets.sort(key=lambda a: a.part_number)

    # Add all selected assets to the document
    for asset in selected_assets:
        file_path = os.path.join(source_path, asset.file_path)

        if not os.path.exists(file_path):
            para = doc.add_paragraph(style='OQB Body Text')
            run = para.add_run(f'[File not found: {asset.file_path}]')
            run.italic = True
            continue

        if asset.file_format == 'IMG':
            try:
                # Open image to get dimensions
                img = Image.open(file_path)
                img_width, img_height = img.size

                # Calculate size for document
                # Max width: 6 inches (to fit in A4 with margins)
                max_width_inches = 6.0
                max_width_pixels = max_width_inches * _DPI

                if img_width > max_width_pixels:
                    doc_width_inches = max_width_inches
                else:
                    doc_width_inches = img_width / _DPI

                # Add picture
                doc.add_picture(file_path, width=Inches(doc_width_inches))

            except Exception as e:
                para = doc.add_paragraph(style='OQB Body Text')
                run = para.add_run(f'[Error loading image: {str(e)}]')
                run.italic = True

        elif asset.file_format == 'MD':
            # Markdown: convert via pandoc -> .docx, then splice into the master doc.
            try:
                _append_md_via_pandoc(doc, file_path)
            except Exception as e:
                para = doc.add_paragraph(style='OQB Body Text')
                run = para.add_run(
                    f'[Error rendering Markdown {asset.file_path}: {e}]'
                )
                run.italic = True

        elif asset.file_format == 'DOC':
            # For Word files, just add a placeholder for now
            para = doc.add_paragraph(style='OQB Body Text')
            run = para.add_run(f'[Word document: {asset.file_path}]')
            run.italic = True