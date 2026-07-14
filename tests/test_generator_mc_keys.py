"""Unit tests for compact MC answer-key generation."""
import unittest
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from flask import Flask

from app import generator


def _question(qid, q_type, answer):
    return SimpleNamespace(qid=qid, q_type=q_type, answer=answer)


class CompactMcOptionTests(unittest.TestCase):
    def test_form_options_are_normalized_and_gated(self):
        options = generator._parse_mc_answer_key_options(
            {
                'compact_mc_answers': 'on',
                'mc_key_layout': 'tabs',
                'mc_key_columns': '99',
                'mc_key_max_rows': '',
                'mc_key_include_seq': 'on',
                'mc_key_range_title': 'on',
            },
            'QUE_THEN_ANS',
            True,
        )
        self.assertTrue(options['enabled'])
        self.assertEqual(options['layout'], 'tabs')
        self.assertEqual(options['columns'], 20)
        self.assertIsNone(options['max_rows'])
        self.assertTrue(options['include_seq'])
        self.assertTrue(options['range_title'])

    def test_options_are_disabled_outside_then_answers_mode(self):
        options = generator._parse_mc_answer_key_options(
            {
                'compact_mc_answers': 'on',
                'mc_key_include_seq': 'on',
                'mc_key_range_title': 'on',
            },
            'QUE_ANS',
            True,
        )
        self.assertFalse(options['enabled'])
        self.assertFalse(options['include_seq'])
        self.assertFalse(options['range_title'])


class CompactMcRunTests(unittest.TestCase):
    def test_eligibility_accepts_short_single_line_text(self):
        self.assertEqual(
            generator._compact_mc_answer_text(
                _question('Q1', 'mc', '  A, C  ')
            ),
            'A, C',
        )
        self.assertIsNone(
            generator._compact_mc_answer_text(
                _question('Q2', 'MC', 'A\nC')
            )
        )
        self.assertIsNone(
            generator._compact_mc_answer_text(
                _question('Q3', 'CQ', 'A')
            )
        )

    def test_partition_preserves_contiguous_runs_and_runtime_numbers(self):
        questions = [
            _question('Q1', 'MC', 'A'),
            _question('Q2', 'MC', 'B'),
            _question('Q3', 'CQ', 'working'),
            _question('Q4', 'MC', ''),
            _question('Q5', 'MC', 'D'),
        ]
        segments = generator._partition_mc_answer_runs(questions, seq_start=21)
        self.assertEqual([segment['type'] for segment in segments],
                         ['mc', 'normal', 'normal', 'mc'])
        self.assertEqual(
            [entry['seq_no'] for entry in segments[0]['entries']],
            [21, 22],
        )
        self.assertEqual(segments[1]['seq_no'], 23)
        self.assertEqual(segments[2]['seq_no'], 24)
        self.assertEqual(segments[3]['entries'][0]['seq_no'], 25)

    def test_max_rows_splits_run_by_capacity(self):
        entries = [{'seq_no': i, 'answer': 'A'} for i in range(1, 13)]
        blocks = generator._split_mc_answer_run(entries, columns=5, max_rows=2)
        self.assertEqual([len(block) for block in blocks], [10, 2])
        self.assertEqual(
            generator._split_mc_answer_run(entries, columns=5, max_rows=None),
            [entries],
        )


class CompactMcDocumentTests(unittest.TestCase):
    def _styled_doc(self):
        doc = Document()
        generator._define_oqb_styles(doc)
        return doc

    def test_table_layout_has_requested_shape_and_partial_row(self):
        doc = self._styled_doc()
        entries = [
            {'seq_no': i, 'answer': chr(64 + i)}
            for i in range(1, 7)
        ]
        generator._add_mc_answer_key_block(
            doc,
            entries,
            {
                'layout': 'table',
                'columns': 5,
                'include_seq': True,
                'range_title': True,
            },
        )
        self.assertEqual(doc.paragraphs[0].text, 'Q1–Q6')
        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(len(table.columns), 5)
        self.assertEqual(table.cell(0, 0).text, '1. A')
        self.assertEqual(table.cell(1, 0).text, '6. F')
        self.assertEqual(table.cell(1, 1).text, '')

    def test_tab_layout_uses_rows_and_can_hide_numbers(self):
        doc = self._styled_doc()
        entries = [
            {'seq_no': i, 'answer': answer}
            for i, answer in enumerate(['A', 'B', 'C', 'D', 'E'], start=7)
        ]
        generator._add_mc_answer_key_block(
            doc,
            entries,
            {
                'layout': 'tabs',
                'columns': 3,
                'include_seq': False,
                'range_title': False,
            },
        )
        self.assertEqual(len(doc.tables), 0)
        self.assertEqual([p.text for p in doc.paragraphs], ['A\tB\tC', 'D\tE'])

    def test_mixed_answers_render_normally_and_suppress_mc_qid(self):
        questions = [
            _question('Q1', 'MC', 'A'),
            _question('Q2', 'CQ', 'full CQ answer'),
            _question('Q3', 'MC', ''),
            _question('Q4', 'MC', 'D'),
        ]
        spacing = {
            'mc': {
                'before_mode': 'lines', 'before_lines': 0,
                'after_mode': 'lines', 'after_lines': 0,
            },
            'cq': {
                'before_mode': 'lines', 'before_lines': 0,
                'after_mode': 'lines', 'after_lines': 0,
            },
        }
        answer_calls = []

        def fake_add(doc, question, asset_type, show_qid, *_args, **_kwargs):
            if asset_type == 'ANS':
                answer_calls.append((question.qid, show_qid))
            doc.add_paragraph(f'{asset_type}:{question.qid}')

        app = Flask(__name__)
        app.config['SOURCE_PATH'] = '.'
        with app.app_context(), patch.object(
            generator, 'add_question_content_to_doc', side_effect=fake_add
        ):
            doc, _ = generator.create_word_document(
                questions,
                'QUE_THEN_ANS',
                spacing,
                show_qid=True,
                show_qid_answer=True,
                show_seq_no=True,
                mc_answer_key_options={
                    'enabled': True,
                    'layout': 'table',
                    'columns': 5,
                    'max_rows': None,
                    'include_seq': True,
                    'range_title': False,
                },
            )

        self.assertEqual(answer_calls, [('Q2', True), ('Q3', False)])
        self.assertEqual(len(doc.tables), 2)


if __name__ == '__main__':
    unittest.main()
